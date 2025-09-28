# gerador_docx.py

import re
from io import BytesIO
from docx import Document

# --- FUNÇÃO ESSENCIAL QUE ESTAVA FALTANDO E FOI REINSERIDA ---
def extrair_nome_aluno(texto_completo_ia):
    """Extrai o nome do aluno da seção '### Nome do Aluno'."""
    match = re.search(r"### Nome do Aluno\s*\n(.*?)\n", texto_completo_ia, re.DOTALL)
    if match:
        nome = match.group(1).strip()
        if nome:
            return nome
    return "Nome_Nao_Identificado"

def extrair_secao(texto_completo, titulo_secao):
    """Extrai o conteúdo de uma seção específica da resposta da IA."""
    padrao = re.compile(f"### {titulo_secao}(.*?)(?=### |$)", re.DOTALL)
    match = padrao.search(texto_completo)
    if match:
        return match.group(1).strip()
    return ""

def buscar_e_substituir(doc_obj, substituicoes):
    """Busca e substitui placeholders em todo o documento (parágrafos e tabelas)."""
    # Para parágrafos
    for p in doc_obj.paragraphs:
        # Cria uma cópia da lista de runs para iterar
        runs = list(p.runs)
        for i, run in enumerate(runs):
            for key, val in substituicoes.items():
                if key in run.text:
                    # Substitui o texto
                    run.text = run.text.replace(key, str(val))

    # Para tabelas (chamada recursiva para cada célula)
    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                buscar_e_substituir(cell, substituicoes)

def deletar_paragrafo(paragraph):
    """Remove um parágrafo do documento."""
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
        p._p = p._element = None

def criar_relatorio_avancado_docx(analise_completa_ia):
    """Cria o relatório .docx final preenchendo o arquivo 'template.docx'."""
    try:
        document = Document('template.docx')

        # Dicionário com todos os placeholders e seus valores
        contexto = {
            '{{NOME_ALUNO}}': extrair_secao(analise_completa_ia, "Nome do Aluno"),
            '{{TEMA}}': extrair_secao(analise_completa_ia, "Tema da Redação"),
            '{{DATA}}': extrair_secao(analise_completa_ia, "Data da Redação"),
            '{{COMENTARIOS}}': extrair_secao(analise_completa_ia, "Comentários Gerais"),
            '{{NOTA_FINAL}}': extrair_secao(analise_completa_ia, "Nota Estimada"),
            '{{ALERTA_ORIGINALIDADE}}': extrair_secao(analise_completa_ia, "Alerta de Originalidade")
        }

        # Extrai as informações de cada competência
        analise_competencias = extrair_secao(analise_completa_ia, "Análise das Competências")
        blocos = re.split(r'(?=\*\*Competência \d)', analise_competencias)
        
        for i in range(1, 6):
            nota_encontrada = ""
            analise_encontrada = ""
            for bloco in blocos:
                if f'**Competência {i}' in bloco:
                    nota = re.search(r'\* \*\*Nota estimada:\*\* (.*?)\n', bloco)
                    analise = re.search(r'\* \*\*Análise.*?\*\* (.*)', bloco, re.DOTALL)
                    if nota: nota_encontrada = nota.group(1).strip()
                    if analise: analise_encontrada = analise.group(1).strip()
                    break
            contexto[f'{{{{NOTA_C{i}}}}}'] = nota_encontrada
            contexto[f'{{{{ANALISE_C{i}}}}}'] = analise_encontrada
        
        # Chama a função para substituir TUDO de uma vez, incluindo as tabelas
        buscar_e_substituir(document, contexto)

        # Lógica para remover a seção de Alerta de Originalidade se estiver vazia
        if not contexto['{{ALERTA_ORIGINALIDADE}}']:
            paragrafos_para_deletar = []
            for p in document.paragraphs:
                # Marca para remoção se encontrar o título ou o placeholder
                if 'Alerta de Originalidade' in p.text or '{{ALERTA_ORIGINALIDADE}}' in p.text:
                    paragrafos_para_deletar.append(p)
            for p in paragrafos_para_deletar:
                deletar_paragrafo(p)
        
        doc_buffer = BytesIO()
        document.save(doc_buffer)
        doc_buffer.seek(0)
        return doc_buffer

    except FileNotFoundError:
        print("❌ ERRO CRÍTICO: O arquivo 'template.docx' não foi encontrado. Certifique-se de que ele está na mesma pasta do projeto.")
        return None
    except Exception as e:
        print(f"❌ Erro ao gerar o arquivo DOCX a partir do template: {e}")
        return None